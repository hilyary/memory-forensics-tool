"""
Report Generator - 生成取证报告
"""

import os
import sys
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


class ReportGenerator:
    """报告生成器类"""

    def __init__(self, image_info: Dict[str, Any]):
        self.image_info = image_info
        self.reports_dir = self._get_user_data_dir() / 'reports'
        self.reports_dir.mkdir(parents=True, exist_ok=True)

    def _get_user_data_dir(self) -> Path:
        """获取用户数据目录"""
        import platform

        system = platform.system()

        if system == 'Darwin':  # macOS
            # macOS: ~/Library/Application Support/LensAnalysis
            app_data = Path.home() / 'Library' / 'Application Support' / 'LensAnalysis'
        elif system == 'Windows':  # Windows
            # Windows: %APPDATA%/LensAnalysis
            app_data = Path(os.environ.get('APPDATA', Path.home() / 'AppData' / 'Roaming')) / 'LensAnalysis'
        else:  # Linux 和其他
            # Linux: ~/.local/share/LensAnalysis (遵循 XDG 规范)
            app_data = Path(os.environ.get('XDG_DATA_HOME', Path.home() / '.local' / 'share')) / 'LensAnalysis'

        return app_data

    def generate_markdown(self, timestamp: str) -> str:
        """生成Markdown格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.md'

        content = self._get_markdown_content()

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"Markdown报告已生成: {report_path}")
        return str(report_path)

    def generate_html(self, timestamp: str) -> str:
        """生成HTML格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.html'

        content = self._get_html_content()

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"HTML报告已生成: {report_path}")
        return str(report_path)

    def generate_docx(self, timestamp: str) -> str:
        """生成Word (docx)格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.docx'

        doc = Document()

        # 添加标题
        title = doc.add_heading('内存取证分析报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加基本信息表格
        doc.add_heading('基本信息', level=1)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        info_data = [
            ('报告生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('镜像文件名', self.image_info.get('name', 'N/A')),
            ('文件大小', self.image_info.get('size', 'N/A')),
            ('文件哈希', self.image_info.get('hash', 'N/A')),
            ('文件路径', self.image_info.get('path', 'N/A'))
        ]

        for i, (label, value) in enumerate(info_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        # 添加支持的分析功能
        doc.add_heading('支持的分析功能', level=1)

        features = [
            '进程分析 - 进程列表 (pslist)、进程树 (pstree)、进程扫描 (psscan)、命令行参数 (cmdline)',
            '网络分析 - 网络连接扫描 (netscan)、网络状态统计 (netstat)',
            '注册表分析 - 注册表配置单元列表 (hivelist)、注册表键值查询 (printkey)',
            '文件系统 - 文件对象扫描 (filescan)、文件列表 (files)',
            '恶意软件分析 - 恶意代码查找 (malfind)、加载模块检测 (ldrmodules)',
            '加密相关 - 密码哈希转储 (hashdump)、LSA密钥转储 (lsadump)',
            'CTF竞赛支持 - 自动Flag搜索、正则表达式匹配、进程内存转储、快捷分析流程'
        ]

        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')

        # 添加技术说明
        doc.add_heading('技术说明', level=1)

        tech_table = doc.add_table(rows=3, cols=2)
        tech_table.style = 'Light Grid Accent 1'

        tech_data = [
            ('分析框架', 'Volatility 3'),
            ('界面技术', 'PyWebView + HTML5'),
            ('编程语言', 'Python 3')
        ]

        for i, (label, value) in enumerate(tech_data):
            tech_table.rows[i].cells[0].text = label
            tech_table.rows[i].cells[1].text = value

        # 添加页脚
        doc.add_paragraph('_' * 50)
        footer = doc.add_paragraph()
        footer.add_run('本报告由 析镜 LensAnalysis 自动生成').italic = True
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(str(report_path))
        logger.info(f"Word报告已生成: {report_path}")
        return str(report_path)

    def generate_markdown_with_data(self, timestamp: str, report_data: Dict[str, Any]) -> str:
        """使用指定数据生成Markdown格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.md'
        content = self._get_markdown_content_with_data(report_data)

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"Markdown报告已生成: {report_path}")
        return str(report_path)

    def generate_html_with_data(self, timestamp: str, report_data: Dict[str, Any]) -> str:
        """使用指定数据生成HTML格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.html'
        content = self._get_html_content_with_data(report_data)

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"HTML报告已生成: {report_path}")
        return str(report_path)

    def generate_docx_with_data(self, timestamp: str, report_data: Dict[str, Any]) -> str:
        """使用指定数据生成Word (docx)格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.docx'
        doc = Document()

        # 添加标题
        title = doc.add_heading('内存取证分析报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加基本信息表格
        doc.add_heading('基本信息', level=1)

        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        info_data = [
            ('报告生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('镜像文件名', self.image_info.get('name', 'N/A')),
            ('文件大小', self.image_info.get('size', 'N/A')),
            ('文件哈希', self.image_info.get('hash', 'N/A')),
            ('操作系统', self.image_info.get('os_type', 'N/A')),
            ('文件路径', self.image_info.get('path', 'N/A'))
        ]

        for i, (label, value) in enumerate(info_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        # 添加分析结果
        plugins = report_data.get('plugins', [])
        if plugins:
            doc.add_heading('分析结果', level=1)
            doc.add_paragraph(f'本次报告包含 {len(plugins)} 个插件的分析结果：')

            for plugin in plugins:
                doc.add_heading(plugin.get('display_name', plugin.get('plugin_id', 'Unknown')), level=2)

                # 插件统计信息
                stats = doc.add_table(rows=3, cols=2)
                stats.style = 'Light Grid Accent 1'
                stats_data = [
                    ('记录数', str(plugin.get('count', 0))),
                    ('执行时间', f"{plugin.get('execution_time', 'N/A')} 秒"),
                    ('执行时间', plugin.get('timestamp', datetime.now().isoformat()))
                ]
                for i, (label, value) in enumerate(stats_data):
                    stats.rows[i].cells[0].text = label
                    stats.rows[i].cells[1].text = str(value)

                # 结果摘要（只显示前5条）
                results = plugin.get('results', [])
                if results:
                    doc.add_paragraph(f'结果摘要（共 {len(results)} 条，显示前 5 条）：')
                    for i, result in enumerate(results[:5]):
                        doc.add_paragraph(f"{i+1}. ", style='List Number')
                        # 将结果转换为可读格式
                        result_text = ' | '.join([f"{k}: {v}" for k, v in result.items() if not k.startswith('_')])
                        doc.add_paragraph(result_text)

        # 添加技术说明
        doc.add_heading('技术说明', level=1)

        tech_table = doc.add_table(rows=3, cols=2)
        tech_table.style = 'Light Grid Accent 1'

        tech_data = [
            ('分析框架', 'Volatility 3'),
            ('界面技术', 'PyWebView + HTML5'),
            ('编程语言', 'Python 3')
        ]

        for i, (label, value) in enumerate(tech_data):
            tech_table.rows[i].cells[0].text = label
            tech_table.rows[i].cells[1].text = value

        # 添加页脚
        doc.add_paragraph('_' * 50)
        footer = doc.add_paragraph()
        footer.add_run('本报告由 析镜 LensAnalysis 自动生成').italic = True
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(str(report_path))
        logger.info(f"Word报告已生成: {report_path}")
        return str(report_path)

    def generate_markdown_from_data(self, plugins_data: list, timestamp: str) -> str:
        """使用提供的插件数据生成Markdown格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.md'

        report_data = {
            'plugins': plugins_data,
            'timestamp': timestamp
        }

        content = self._get_markdown_content_with_data(report_data)

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"Markdown报告已生成（使用提供的数据）: {report_path}")
        return str(report_path)

    def generate_html_from_data(self, plugins_data: list, timestamp: str) -> str:
        """使用提供的插件数据生成HTML格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.html'

        report_data = {
            'plugins': plugins_data,
            'timestamp': timestamp
        }

        content = self._get_html_content_with_data(report_data)

        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)

        logger.info(f"HTML报告已生成（使用提供的数据）: {report_path}")
        return str(report_path)

    def generate_docx_from_data(self, plugins_data: list, timestamp: str) -> str:
        """使用提供的插件数据生成Word (docx)格式报告"""
        report_path = self.reports_dir / f'forensics_report_{timestamp}.docx'

        # 计算总记录数
        total_records = sum(len(p.get('results', [])) for p in plugins_data)

        doc = Document()

        # 添加标题
        title = doc.add_heading('内存取证分析报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加基本信息表格
        doc.add_heading('基本信息', level=1)

        # 获取符号表文件名
        symbol_file = self.image_info.get('symbol_file', '未安装')

        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'

        info_data = [
            ('报告生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('镜像文件名', self.image_info.get('name', 'N/A')),
            ('文件大小', self.image_info.get('size', 'N/A')),
            ('文件哈希', self.image_info.get('hash', 'N/A')),
            ('操作系统', self.image_info.get('os_type', 'N/A')),
            ('文件路径', self.image_info.get('path', 'N/A')),
            ('符号表', symbol_file)
        ]

        for i, (label, value) in enumerate(info_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        # 添加分析结果
        if plugins_data:
            doc.add_heading('分析结果', level=1)
            doc.add_paragraph(f'本次报告包含 {len(plugins_data)} 个插件的分析结果（共 {total_records} 条记录）：')

            for plugin in plugins_data:
                display_name = plugin.get('display_name', plugin.get('plugin_id', 'Unknown'))
                doc.add_heading(display_name, level=2)

                # 插件统计信息
                timestamp_str = plugin.get('timestamp', '')
                results = plugin.get('results', [])

                stats = doc.add_table(rows=2, cols=2)
                stats.style = 'Light Grid Accent 1'
                stats_data = [
                    ('记录数', str(len(results))),
                    ('执行时间', timestamp_str)
                ]
                for i, (label, value) in enumerate(stats_data):
                    stats.rows[i].cells[0].text = label
                    stats.rows[i].cells[1].text = str(value)

                # 添加完整结果表格
                if results:
                    doc.add_paragraph(f'完整结果（共 {len(results)} 条）：')

                    # 获取表头
                    first_result = results[0]
                    headers = [k for k in first_result.keys() if not k.startswith('_')]

                    # 创建结果表格
                    result_table = doc.add_table(rows=len(results) + 1, cols=len(headers))
                    result_table.style = 'Light List Accent 1'

                    # 表头 - 批量设置
                    header_row = result_table.rows[0]
                    for col_idx, header in enumerate(headers):
                        header_row.cells[col_idx].text = header
                        # 更高效的方式设置表头加粗
                        for paragraph in header_row.cells[col_idx].paragraphs:
                            if paragraph.runs:
                                paragraph.runs[0].font.bold = True
                            else:
                                paragraph.add_run(header).bold = True

                    # 数据行 - 直接批量写入
                    table_rows = result_table.rows[1:]
                    for row_idx, result in enumerate(results):
                        row_cells = table_rows[row_idx].cells
                        for col_idx, header in enumerate(headers):
                            val = str(result.get(header, ''))[:100]  # 限制长度
                            row_cells[col_idx].text = val

        # 添加技术说明
        doc.add_heading('技术说明', level=1)

        tech_table = doc.add_table(rows=3, cols=2)
        tech_table.style = 'Light Grid Accent 1'

        tech_data = [
            ('分析框架', 'Volatility 3'),
            ('界面技术', 'PyWebView + HTML5'),
            ('编程语言', 'Python 3')
        ]

        for i, (label, value) in enumerate(tech_data):
            tech_table.rows[i].cells[0].text = label
            tech_table.rows[i].cells[1].text = value

        # 添加页脚
        doc.add_paragraph('_' * 50)
        footer = doc.add_paragraph()
        footer.add_run('本报告由 析镜 LensAnalysis 自动生成').italic = True
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(str(report_path))
        logger.info(f"Word报告已生成（使用提供的数据）: {report_path}")
        return str(report_path)

    def _get_markdown_content_with_data(self, report_data: Dict[str, Any]) -> str:
        """获取带数据的Markdown报告内容"""
        plugins = report_data.get('plugins', [])
        plugin_summary = '\n'.join([f"- **{p.get('display_name', p.get('plugin_id', 'Unknown'))}** - {p.get('count', 0)} 条结果" for p in plugins])

        # 获取符号表文件名
        symbol_file = self.image_info.get('symbol_file', '未安装')

        # 构建每个插件的详细结果
        plugins_detail = ""
        for plugin in plugins:
            display_name = plugin.get('display_name', plugin.get('plugin_id', 'Unknown'))
            results = plugin.get('results', [])
            timestamp = plugin.get('timestamp', '')

            plugins_detail += f"""
### {display_name}

**执行时间**: {timestamp}
**记录数**: {len(results)}

"""

            # 添加完整结果表格
            if results:
                # 获取表头
                first_result = results[0]
                headers = [k for k in first_result.keys() if not k.startswith('_')]

                # 表头
                plugins_detail += "| " + " | ".join(headers) + " |\n"
                plugins_detail += "| " + " | ".join(["---"] * len(headers)) + " |\n"

                # 数据行
                for result in results:
                    row_values = [str(result.get(h, ''))[:100] for h in headers]  # 限制每列长度
                    plugins_detail += "| " + " | ".join(row_values) + " |\n"

            plugins_detail += "\n---\n"

        return f"""# 内存取证分析报告

> 本报告使用「析镜 LensAnalysis」生成

## 基本信息

- **报告生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- **镜像文件名**: {self.image_info.get('name', 'N/A')}
- **文件大小**: {self.image_info.get('size', 'N/A')}
- **文件哈希**: {self.image_info.get('hash', 'N/A')}
- **操作系统**: {self.image_info.get('os_type', 'N/A')}
- **符号表**: {symbol_file}

## 分析结果

本次报告包含 **{len(plugins)}** 个插件的分析结果：

{plugin_summary}

{plugins_detail}

---

*本报告由 析镜 LensAnalysis 自动生成*
"""

    def _get_html_content_with_data(self, report_data: Dict[str, Any]) -> str:
        """获取带数据的HTML报告内容"""
        plugins = report_data.get('plugins', [])
        plugin_list_html = '\n'.join([f"<li><strong>{p.get('display_name', p.get('plugin_id', 'Unknown'))}</strong> - {p.get('count', 0)} 条结果</li>" for p in plugins])

        # 获取符号表文件名
        symbol_file = self.image_info.get('symbol_file', '未安装')

        # 构建每个插件的详细结果
        plugins_detail_html = ""
        for plugin in plugins:
            display_name = plugin.get('display_name', plugin.get('plugin_id', 'Unknown'))
            results = plugin.get('results', [])
            timestamp = plugin.get('timestamp', '')

            plugins_detail_html += f"""
        <h3>{display_name}</h3>
        <p><strong>执行时间</strong>: {timestamp} | <strong>记录数</strong>: {len(results)}</p>
"""

            if results:
                plugins_detail_html += '        <table class="result-table">\n'
                # 获取表头
                first_result = results[0]
                headers = [k for k in first_result.keys() if not k.startswith('_')]

                # 表头
                plugins_detail_html += '            <thead><tr>\n'
                for h in headers:
                    plugins_detail_html += f'                <th>{h}</th>\n'
                plugins_detail_html += '            </tr></thead>\n'
                plugins_detail_html += '            <tbody>\n'

                # 数据行
                for result in results:
                    plugins_detail_html += '            <tr>\n'
                    for h in headers:
                        val = str(result.get(h, ''))[:100]
                        plugins_detail_html += f'                <td>{val}</td>\n'
                    plugins_detail_html += '            </tr>\n'

                plugins_detail_html += '            </tbody>\n        </table>\n'

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>内存取证分析报告</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Oxygen', 'Ubuntu', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f5f5f5;
            padding: 20px;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #0a0e27;
            border-bottom: 3px solid #00d4ff;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }}
        h2 {{
            color: #111830;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        h3 {{
            color: #333;
            margin-top: 25px;
            margin-bottom: 10px;
            font-size: 18px;
        }}
        .info-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        .info-table tr {{
            border-bottom: 1px solid #eee;
        }}
        .info-table td {{
            padding: 12px 8px;
        }}
        .info-table td:first-child {{
            font-weight: 600;
            color: #666;
            width: 150px;
        }}
        .result-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
            font-size: 13px;
        }}
        .result-table th,
        .result-table td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        .result-table th {{
            background: #f8f9fa;
            font-weight: 600;
        }}
        .result-table tr:nth-child(even) {{
            background: #f8f9fa;
        }}
        .plugin-list {{
            list-style: none;
            padding: 0;
        }}
        .plugin-list li {{
            padding: 10px 0;
            padding-left: 24px;
            position: relative;
        }}
        .plugin-list li:before {{
            content: "✓";
            position: absolute;
            left: 0;
            color: #00d4ff;
            font-weight: bold;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #eee;
            text-align: center;
            color: #999;
            font-size: 14px;
        }}
        code {{
            background: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 12px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>内存取证分析报告</h1>

        <h2>基本信息</h2>
        <table class="info-table">
            <tr>
                <td>报告生成时间</td>
                <td>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</td>
            </tr>
            <tr>
                <td>镜像文件名</td>
                <td>{self.image_info.get('name', 'N/A')}</td>
            </tr>
            <tr>
                <td>文件大小</td>
                <td>{self.image_info.get('size', 'N/A')}</td>
            </tr>
            <tr>
                <td>文件哈希</td>
                <td>{self.image_info.get('hash', 'N/A')}</td>
            </tr>
            <tr>
                <td>操作系统</td>
                <td>{self.image_info.get('os_type', 'N/A')}</td>
            </tr>
            <tr>
                <td>符号表</td>
                <td>{symbol_file}</td>
            </tr>
        </table>

        <h2>分析结果</h2>
        <p>本次报告包含 <strong>{len(plugins)}</strong> 个插件的分析结果：</p>
        <ul class="plugin-list">
            {plugin_list_html}
        </ul>

        <hr style="margin: 30px 0; border: none; border-top: 1px solid #eee;">

        <h2>详细结果</h2>
        {plugins_detail_html}

        <div class="footer">
            本报告由 析镜 LensAnalysis 自动生成
        </div>
    </div>
</body>
</html>
"""

    def _get_markdown_content(self) -> str:
        return f"""# 内存取证分析报告

> 本报告使用「析镜 LensAnalysis」生成

## 基本信息

- **报告生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- **镜像文件名**: {self.image_info.get('name', 'N/A')}
- **文件大小**: {self.image_info.get('size', 'N/A')}
- **文件哈希**: {self.image_info.get('hash', 'N/A')}

## 分析摘要

本报告使用基于 Volatility 3 框架的图形化内存取证工具生成。

### 支持的分析功能

1. **进程分析**
   - 进程列表 (pslist)
   - 进程树 (pstree)
   - 进程扫描 (psscan)
   - 命令行参数 (cmdline)

2. **网络分析**
   - 网络连接扫描 (netscan)
   - 网络状态统计 (netstat)

3. **注册表分析**
   - 注册表配置单元列表 (hivelist)
   - 注册表键值查询 (printkey)

4. **文件系统**
   - 文件对象扫描 (filescan)
   - 文件列表 (files)

5. **恶意软件分析**
   - 恶意代码查找 (malfind)
   - 加载模块检测 (ldrmodules)

6. **加密相关**
   - 密码哈希转储 (hashdump)
   - LSA密钥转储 (lsadump)

## CTF 竞赛支持

本工具特别针对 CTF 竞赛场景进行了优化：

- 自动 Flag 搜索功能
- 正则表达式匹配
- 进程内存转储
- 快捷分析流程

## 技术说明

- **分析框架**: Volatility 3
- **界面技术**: PyWebView + HTML5
- **编程语言**: Python 3

---
*本报告由 析镜 LensAnalysis 自动生成*
"""

    def _get_html_content(self) -> str:
        """获取HTML报告内容"""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>内存取证分析报告</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Oxygen', 'Ubuntu', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f5f5f5;
            padding: 20px;
        }}

        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}

        h1 {{
            color: #0a0e27;
            border-bottom: 3px solid #00d4ff;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }}

        h2 {{
            color: #111830;
            margin-top: 30px;
            margin-bottom: 15px;
        }}

        .info-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}

        .info-table tr {{
            border-bottom: 1px solid #eee;
        }}

        .info-table td {{
            padding: 12px 8px;
        }}

        .info-table td:first-child {{
            font-weight: 600;
            color: #666;
            width: 150px;
        }}

        .feature-list {{
            list-style: none;
            padding: 0;
        }}

        .feature-list li {{
            padding: 10px 0;
            padding-left: 24px;
            position: relative;
        }}

        .feature-list li:before {{
            content: "✓";
            position: absolute;
            left: 0;
            color: #00d4ff;
            font-weight: bold;
        }}

        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #eee;
            text-align: center;
            color: #999;
            font-size: 14px;
        }}

        .tag {{
            display: inline-block;
            padding: 4px 12px;
            background: #e3f2fd;
            color: #1976d2;
            border-radius: 4px;
            font-size: 12px;
            margin-right: 8px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>内存取证分析报告</h1>

        <h2>基本信息</h2>
        <table class="info-table">
            <tr>
                <td>报告生成时间</td>
                <td>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</td>
            </tr>
            <tr>
                <td>镜像文件名</td>
                <td>{self.image_info.get('name', 'N/A')}</td>
            </tr>
            <tr>
                <td>文件大小</td>
                <td>{self.image_info.get('size', 'N/A')}</td>
            </tr>
            <tr>
                <td>文件哈希</td>
                <td>{self.image_info.get('hash', 'N/A')}</td>
            </tr>
        </table>

        <h2>支持的分析功能</h2>
        <ul class="feature-list">
            <li>进程分析 (进程列表、进程树、进程扫描、命令行参数)</li>
            <li>网络分析 (网络连接扫描、网络状态统计)</li>
            <li>注册表分析 (配置单元列表、键值查询)</li>
            <li>文件系统 (文件对象扫描、文件列表)</li>
            <li>恶意软件分析 (恶意代码查找、模块检测)</li>
            <li>加密相关 (密码哈希转储、LSA密钥转储)</li>
        </ul>

        <h2>CTF 竞赛支持</h2>
        <p>本工具特别针对 CTF 竞赛场景进行了优化：</p>
        <div style="margin: 15px 0;">
            <span class="tag">自动 Flag 搜索</span>
            <span class="tag">正则表达式匹配</span>
            <span class="tag">进程内存转储</span>
            <span class="tag">快捷分析流程</span>
        </div>

        <h2>技术说明</h2>
        <table class="info-table">
            <tr>
                <td>分析框架</td>
                <td>Volatility 3</td>
            </tr>
            <tr>
                <td>界面技术</td>
                <td>PyWebView + HTML5</td>
            </tr>
            <tr>
                <td>编程语言</td>
                <td>Python 3</td>
            </tr>
        </table>

        <div class="footer">
            本报告由 析镜 LensAnalysis 自动生成
        </div>
    </div>
</body>
</html>
"""
